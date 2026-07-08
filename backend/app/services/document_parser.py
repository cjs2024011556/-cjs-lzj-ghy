"""
文档解析服务
支持: PDF（结构化）/ Word / Markdown / TXT / 图片（OCR）

聚群 A 升级：
- PDF 走 PdfplumberPdfLoader → StructureDetector 流水线
- 每 chunk 携带: page_number / page_end / chapter / section_title /
                section_type ('text'|'table'|'heading') / section_level /
                doc_id / keywords

聚群 B 升级：
- 扫描页 / 视觉重页走 VL 图像理解（image_description + image_facts）
- 由 parse_pdf_with_vl() 异步方法触发
- 默认关闭（保持现有流程稳定）；可注入 mock describer 测试
"""
import re
import uuid
from pathlib import Path
from typing import List, Dict, Any, Optional
from loguru import logger

from app.core.config import settings


class DocumentChunk:
    """文档分块"""
    def __init__(self, content: str, metadata: Dict[str, Any] = None):
        self.content = content
        self.metadata = metadata or {}
        self.chunk_id = ""

    def __repr__(self):
        return (
            f"<DocumentChunk id={self.chunk_id} "
            f"page={self.metadata.get('page_number', '?')}/"
            f"{self.metadata.get('page_end', '?')} "
            f"type={self.metadata.get('section_type', '?')}>"
        )


class DocumentParser:
    """文档解析器 — 兼容旧接口

    流程:
    - PDF  : PdfplumberPdfLoader → StructureDetector → SemanticChunker（带元数据）
    - Word : 旧版段落+表格抽取（无 page 元数据）
    - MD   : 直接读 UTF-8（无 page 元数据，但有 ## 章节）
    - TXT  : 直接读 UTF-8
    - IMG  : OCR（多模态 LLM）
    """

    DEFAULT_CHUNK_SIZE = 500
    DEFAULT_CHUNK_OVERLAP = 50

    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    # ============================================================
    # 主入口
    # ============================================================
    def parse(self, file_path: str) -> List[DocumentChunk]:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        suffix = path.suffix.lower()

        # —— PDF 走结构化管线 ——
        if suffix == ".pdf":
            return self._parse_pdf_structured(path)

        # —— 其他格式走旧路径 ——
        if suffix in (".docx", ".doc"):
            text = self._parse_word(path)
        elif suffix in (".md", ".markdown"):
            text = self._parse_markdown(path)
        elif suffix == ".txt":
            text = path.read_text(encoding="utf-8")
        elif suffix in (".jpg", ".jpeg", ".png", ".bmp", ".webp"):
            text = self._parse_image(path)
        else:
            raise ValueError(f"不支持的文件类型: {suffix}")

        logger.info(f"解析完成: {path.name}, 文本长度: {len(text)}")
        chunks = self._split_into_chunks(text, source=str(path))
        logger.info(f"切分完成: {len(chunks)} 个块")
        return chunks

    # ============================================================
    # 聚群 B: PDF + VL 异步入口
    # ============================================================
    async def parse_pdf_with_vl(
        self,
        file_path: str,
        enable_vl: bool = True,
        describer: Optional[Any] = None,
        renderer: Optional[Any] = None,
        vl_dpi: int = 150,
        trigger_on_scanned: bool = True,
        trigger_on_table_pages: bool = False,
    ) -> List[DocumentChunk]:
        """PDF 结构化解析 + VL 图像增强（聚群 B）

        流程：
        1. 走 _parse_pdf_structured（同聚群 A）
        2. 对每页渲染 PNG（PageRenderer）
        3. 调 VL 提取 image_description + image_facts
        4. 把 VL 结果回填到该页所有 chunk 的 metadata

        Args:
            file_path: PDF 路径
            enable_vl: 是否启用 VL（默认 True）
            describer: 注入的 ImageDescriber（None 用默认；测试可传 MockImageDescriber）
            renderer: 注入的 PageRenderer（None 用默认）
            vl_dpi: 渲染 DPI（150 足够，200 更高精度）
            trigger_on_scanned: 扫描页（text<50 chars）触发 VL
            trigger_on_table_pages: 含表格的页触发 VL（更慢但更全）

        Returns:
            List[DocumentChunk]（metadata 含 image_description / image_facts）
        """
        path = Path(file_path)
        chunks = self._parse_pdf_structured(path)
        if not enable_vl or not chunks:
            return chunks
        return await self._enrich_chunks_with_vl(
            chunks, path,
            describer=describer,
            renderer=renderer,
            vl_dpi=vl_dpi,
            trigger_on_scanned=trigger_on_scanned,
            trigger_on_table_pages=trigger_on_table_pages,
        )

    async def _enrich_chunks_with_vl(
        self,
        chunks: List[DocumentChunk],
        pdf_path: Path,
        describer: Optional[Any] = None,
        renderer: Optional[Any] = None,
        vl_dpi: int = 150,
        trigger_on_scanned: bool = True,
        trigger_on_table_pages: bool = False,
    ) -> List[DocumentChunk]:
        """为 chunks 添加 image_description / image_facts 字段（聚群 B）"""
        # 延迟导入（避免非 PDF 路径触发）
        from app.services.page_renderer import PageRenderer
        from app.services.image_describer import ImageDescriber

        # 找出需要 VL 的页（按 page 聚合）
        pages_to_describe: Dict[int, List[DocumentChunk]] = {}
        for c in chunks:
            pn = c.metadata.get("page_number", 0)
            if pn <= 0:
                continue
            # 触发判断（按 page 聚合：同一页所有 chunk 共享一次 VL 调用）
            should_trigger = False
            if trigger_on_scanned and c.metadata.get("is_likely_scanned"):
                should_trigger = True
            elif trigger_on_table_pages and c.metadata.get("section_type") == "table":
                should_trigger = True
            else:
                # 兜底启发：内容极短（很可能是图片占主导的页）
                if len(c.content.strip()) < 30:
                    should_trigger = True
            if should_trigger:
                pages_to_describe.setdefault(pn, []).append(c)

        if not pages_to_describe:
            logger.debug("VL 跳过：未检测到需要视觉理解的页")
            return chunks

        renderer = renderer or PageRenderer(dpi=vl_dpi)
        describer = describer or ImageDescriber()

        logger.info(f"🖼️ 聚群 B VL 增强：{len(pages_to_describe)} 页待处理")
        for page_num in sorted(pages_to_describe.keys()):
            page_chunks = pages_to_describe[page_num]
            try:
                png = renderer.render_page(pdf_path, page_num, dpi=vl_dpi)
                desc = await describer.describe(png, page_number=page_num)
                if desc.is_empty():
                    continue
                for c in page_chunks:
                    c.metadata["image_description"] = desc.description[:2000]
                    c.metadata["image_facts"] = ",".join(desc.facts)[:1000]
            except Exception as e:
                logger.warning(f"⚠️ Page {page_num} VL 失败: {e}")

        enriched = sum(
            1 for c in chunks if c.metadata.get("image_description")
        )
        logger.info(
            f"🖼️ 聚群 B 完成：{enriched}/{len(chunks)} chunks 含 image_description"
        )
        return chunks

    # ============================================================
    # PDF 结构化管线（聚群 A 核心）
    # ============================================================
    def _parse_pdf_structured(self, path: Path) -> List[DocumentChunk]:
        """PDF → List[DocumentChunk with rich metadata]

        流程：
        1. PdfplumberPdfLoader.load() → List[PageDoc]
        2. StructureDetector.detect() → List[Section]
        3. 对每个 Section 转 chunk
           - heading-only（H1/H2/H3，无 body）→ 转 1 个 heading chunk
           - 有 body_text → 按段落切 + 长度控制
           - 有 tables → 每张表转 1 个 markdown chunk（section_type=table）
        """
        # 延迟导入避免 docx 解析时硬错误
        from app.services.pdf_loader import PdfplumberPdfLoader
        from app.services.structure_detector import StructureDetector

        # 1. 加载 PDF
        try:
            pages = PdfplumberPdfLoader().load(path)
        except ImportError:
            # pdfplumber 不在：回退到 pypdf 简易模式
            logger.warning("pdfplumber 不可用，回退到 pypdf（无结构信息）")
            return self._parse_pdf_fallback_pypdf(path)

        if not pages:
            return []

        # 2. 识别结构
        try:
            sections = StructureDetector().detect(pages)
        except Exception as e:
            logger.warning(f"StructureDetector 失败，回退纯文本流: {e}")
            # 退化：整页当一段
            all_text = "\n\n".join(p.body_text for p in pages if p.body_text)
            return self._split_into_chunks(all_text, source=str(path))

        # 3. doc_id（同一 PDF 各 chunk 共享）
        doc_id = uuid.uuid5(uuid.NAMESPACE_URL, str(path.absolute())).hex[:16]
        source = str(path)
        source_name = path.name

        chunks: List[DocumentChunk] = []
        chunk_idx = 0

        for sec in sections:
            # 3.1 章节归属元数据
            chapter = sec.chapter or source_name
            section_title = sec.title if sec.title != "(前言)" else ""
            page_start = sec.page_start
            page_end = sec.page_end

            # 3.2 heading-only：单独 1 个 chunk（作锚点）
            if sec.level >= 1 and not sec.has_body and sec.table_count == 0:
                content = sec.title
                chunks.append(self._make_chunk(
                    content=content,
                    source=source,
                    idx=chunk_idx,
                    page_number=page_start,
                    page_end=page_end,
                    chapter=chapter,
                    section_title=section_title,
                    section_type="heading",
                    section_level=sec.level,
                    doc_id=doc_id,
                ))
                chunk_idx += 1
                continue

            # 3.3 表格：每张表 1 个 chunk
            for page_num, table in sec.tables:
                md = PdfplumberPdfLoader.table_to_markdown(table)
                if not md.strip():
                    continue
                # 在内容前加章节标题作为上下文
                content_parts = []
                if section_title:
                    content_parts.append(f"《{section_title}》\n")
                content_parts.append(md)
                content = "\n".join(content_parts)
                chunks.append(self._make_chunk(
                    content=content,
                    source=source,
                    idx=chunk_idx,
                    page_number=page_num,
                    page_end=page_num,
                    chapter=chapter,
                    section_title=section_title,
                    section_type="table",
                    section_level=sec.level,
                    doc_id=doc_id,
                ))
                chunk_idx += 1

            # 3.4 正文：按段落切，长度控制
            if sec.has_body:
                body_chunks = self._split_text_with_section_meta(
                    text=sec.body_text,
                    source=source,
                    source_name=source_name,
                    page_start=page_start,
                    page_end=page_end,
                    chapter=chapter,
                    section_title=section_title,
                    section_level=sec.level,
                    doc_id=doc_id,
                    start_idx=chunk_idx,
                )
                chunks.extend(body_chunks)
                chunk_idx += len(body_chunks)

        logger.info(
            f"📦 PDF 结构化解析 {path.name}: {len(chunks)} chunks "
            f"({sum(1 for c in chunks if c.metadata.get('section_type')=='heading')} 标题 / "
            f"{sum(1 for c in chunks if c.metadata.get('section_type')=='table')} 表 / "
            f"{sum(1 for c in chunks if c.metadata.get('section_type')=='text')} 正文)"
        )
        return chunks

    # ============================================================
    # 段落切分（带 Section 元数据）
    # ============================================================
    def _split_text_with_section_meta(
        self,
        text: str,
        source: str,
        source_name: str,
        page_start: int,
        page_end: int,
        chapter: str,
        section_title: str,
        section_level: int,
        doc_id: str,
        start_idx: int,
    ) -> List[DocumentChunk]:
        """把 Section 的 body_text 按段落 + 长度切分（保留章节元数据）"""
        paragraphs = [p.strip() for p in re.split(r'\n\s*\n', text) if p.strip()]
        chunks: List[DocumentChunk] = []
        current = ""
        idx = start_idx

        for para in paragraphs:
            if len(para) > self.chunk_size:
                # 单段超长：按句子切
                if current:
                    chunks.append(self._make_chunk(
                        content=current,
                        source=source,
                        idx=idx,
                        page_number=page_start,
                        page_end=page_end,
                        chapter=chapter,
                        section_title=section_title,
                        section_type="text",
                        section_level=section_level,
                        doc_id=doc_id,
                    ))
                    idx += 1
                    current = ""
                sentences = re.split(r'([。！？；\.!?;])', para)
                buf = ""
                for s in sentences:
                    buf += s
                    if len(buf) >= self.chunk_size:
                        chunks.append(self._make_chunk(
                            content=buf,
                            source=source,
                            idx=idx,
                            page_number=page_start,
                            page_end=page_end,
                            chapter=chapter,
                            section_title=section_title,
                            section_type="text",
                            section_level=section_level,
                            doc_id=doc_id,
                        ))
                        idx += 1
                        buf = ""
                if buf:
                    current = buf
                continue

            # 累计拼接
            if len(current) + len(para) > self.chunk_size and current:
                chunks.append(self._make_chunk(
                    content=current,
                    source=source,
                    idx=idx,
                    page_number=page_start,
                    page_end=page_end,
                    chapter=chapter,
                    section_title=section_title,
                    section_type="text",
                    section_level=section_level,
                    doc_id=doc_id,
                ))
                idx += 1
                # overlap
                if self.chunk_overlap > 0 and len(current) > self.chunk_overlap:
                    current = current[-self.chunk_overlap:] + "\n\n" + para
                else:
                    current = para
            else:
                current = current + "\n\n" + para if current else para

        if current:
            chunks.append(self._make_chunk(
                content=current,
                source=source,
                idx=idx,
                page_number=page_start,
                page_end=page_end,
                chapter=chapter,
                section_title=section_title,
                section_type="text",
                section_level=section_level,
                doc_id=doc_id,
            ))
            idx += 1

        return chunks

    # ============================================================
    # 兜底：pdfplumber 不可用，用 pypdf 走老路径
    # ============================================================
    def _parse_pdf_fallback_pypdf(self, path: Path) -> List[DocumentChunk]:
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ImportError("请安装 pdfplumber 或 pypdf: pip install pdfplumber pypdf")
        reader = PdfReader(str(path))
        full_text_parts = []
        for i, p in enumerate(reader.pages, start=1):
            t = p.extract_text() or ""
            full_text_parts.append(t)
        return self._split_into_chunks("\n\n".join(full_text_parts), source=str(path))

    # ============================================================
    # 旧版：纯文本/Word/MD/TXT/IMG
    # ============================================================
    def _parse_word(self, path: Path) -> str:
        try:
            from docx import Document
        except ImportError:
            raise ImportError("请安装 python-docx: pip install python-docx")
        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        return "\n\n".join(paragraphs)

    def _parse_markdown(self, path: Path) -> str:
        return path.read_text(encoding="utf-8")

    def _parse_image(self, path: Path) -> str:
        from app.llm.factory import get_model_adapter
        import asyncio

        async def _ocr():
            adapter = get_model_adapter()
            return await adapter.parse_document(str(path))

        try:
            loop = asyncio.get_event_loop()
            if loop.is_running():
                return asyncio.create_task(_ocr())
            else:
                return loop.run_until_complete(_ocr())
        except RuntimeError:
            return asyncio.run(_ocr())

    def _split_into_chunks(self, text: str, source: str) -> List[DocumentChunk]:
        """旧版段落切分（用于非 PDF 格式，保留向后兼容）"""
        paragraphs = [p.strip() for p in re.split(r'\n\s*\n', text) if p.strip()]
        chunks: List[DocumentChunk] = []
        current = ""
        chunk_idx = 0
        for para in paragraphs:
            if len(para) > self.chunk_size:
                if current:
                    chunks.append(self._make_chunk(current, source, chunk_idx))
                    chunk_idx += 1
                    current = ""
                sentences = re.split(r'([。！？；\.!?;])', para)
                buf = ""
                for s in sentences:
                    buf += s
                    if len(buf) >= self.chunk_size:
                        chunks.append(self._make_chunk(buf, source, chunk_idx))
                        chunk_idx += 1
                        buf = ""
                if buf:
                    current = buf
                continue
            if len(current) + len(para) > self.chunk_size and current:
                chunks.append(self._make_chunk(current, source, chunk_idx))
                chunk_idx += 1
                if self.chunk_overlap > 0 and len(current) > self.chunk_overlap:
                    current = current[-self.chunk_overlap:] + "\n\n" + para
                else:
                    current = para
            else:
                current = current + "\n\n" + para if current else para
        if current:
            chunks.append(self._make_chunk(current, source, chunk_idx))
        return chunks

    # ============================================================
    # 构造 chunk
    # ============================================================
    def _make_chunk(
        self,
        content: str,
        source: str,
        idx: int,
        page_number: int = 0,
        page_end: int = 0,
        chapter: str = "",
        section_title: str = "",
        section_type: str = "text",  # text / table / heading
        section_level: int = 0,
        doc_id: str = "",
        is_likely_scanned: bool = False,
        image_description: str = "",
        image_facts: str = "",
    ) -> DocumentChunk:
        return DocumentChunk(
            content=content.strip(),
            metadata={
                "source": source,
                "chunk_index": idx,
                "length": len(content),
                "page_number": page_number,
                "page_end": page_end,
                "chapter": chapter,
                "section_title": section_title,
                "section_type": section_type,
                "section_level": section_level,
                "doc_id": doc_id,
                "keywords": self._extract_keywords(content),
                # 聚群 B 字段
                "is_likely_scanned": is_likely_scanned,
                "image_description": image_description,
                "image_facts": image_facts,
            }
        )

    @staticmethod
    def _extract_keywords(content: str, topk: int = 8) -> str:
        """轻量关键字提取（jieba，可选）"""
        if not content or not content.strip():
            return ""
        try:
            import jieba
            import jieba.analyse
            tags = jieba.analyse.extract_tags(content, topK=topk)
            return ",".join(tags)
        except Exception:
            # jieba 不在 → 退化到字符级首尾
            cleaned = re.sub(r'\s+', ' ', content).strip()
            return cleaned[:60]


def parse_document(file_path: str) -> List[DocumentChunk]:
    """便捷函数"""
    return DocumentParser().parse(file_path)
