"""
pytest 共享 fixtures（P0: 端到端测试公共基础设施）

策略：
- 复用 FastAPI TestClient（与已有 test_api.py 一致，避免 httpx 异步副作用）
- 用 tmp_path 隔离上传文件，避免污染真实 MANUALS_DIR
- 通过 monkey-patch 模块级常量 MANUALS_DIR，让 upload/list/delete 走临时目录
- 测试结束后自动清理，不留垃圾文件
"""
import io
import shutil
from pathlib import Path
from typing import List

import pytest
from fastapi.testclient import TestClient


# ============================================================
# 路径与环境
# ============================================================

# 把 backend 加入 sys.path，确保 `from app.xxx import ...` 可工作
import sys

_BACKEND_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(_BACKEND_ROOT))


@pytest.fixture(scope="session")
def project_root() -> Path:
    """backend/ 目录"""
    return _BACKEND_ROOT


@pytest.fixture(scope="session")
def real_manuals_dir(project_root) -> Path:
    """真实手册目录（用于读取内置数据）"""
    return project_root / "data" / "raw" / "manuals"


# ============================================================
# FastAPI app + TestClient
# ============================================================

@pytest.fixture(scope="session")
def app():
    """共享的 FastAPI app 实例（session 级别：启动一次）"""
    from app.main import app as fastapi_app
    return fastapi_app


@pytest.fixture()
def client(app):
    """FastAPI TestClient（每个测试一个新实例，自动管理 lifespan）"""
    with TestClient(app) as c:
        yield c


# ============================================================
# 临时 manuals 目录（隔离测试文件）
# ============================================================

@pytest.fixture()
def temp_manuals_dir(tmp_path, monkeypatch):
    """给每个测试一个全新 manuals 目录，monkey-patch 模块级常量"""
    test_dir = tmp_path / "manuals"
    test_dir.mkdir(parents=True, exist_ok=True)

    # knowledge.py 在 import 时已绑定 MANUALS_DIR（在模块级 from app.core.config import ...）
    # → 必须同时 patch 模块级属性
    from app.api.v1 import knowledge as knowledge_mod
    monkeypatch.setattr(knowledge_mod, "MANUALS_DIR", test_dir)

    # chat.py 的 _keyword_fallback 在函数内部做 `from app.core.config import MANUALS_DIR`
    #   所以要在 app.core.config 模块上 patch（这样函数内 import 时拿到新值）
    from app.core import config as config_mod
    monkeypatch.setattr(config_mod, "MANUALS_DIR", test_dir)

    yield test_dir

    # 测试结束自动清理（tmp_path 也会清，但显式更安全）
    if test_dir.exists():
        shutil.rmtree(test_dir, ignore_errors=True)


# ============================================================
# 测试数据工厂
# ============================================================

@pytest.fixture()
def sample_md_text() -> str:
    """可被关键词命中的样例 MD 文本（关于焊接机器人）"""
    return """# 焊接机器人故障处理手册

## 焊接飞溅过大

焊接飞溅是常见焊接缺陷，主要原因及解决方案如下：

### 原因一：电流过大
焊接电流超过工艺标准时，会导致熔池温度过高，金属飞溅加剧。
**处理方法**：调低电流 10-15%，观察飞溅变化。

### 原因二：保护气体流量不足
氩气流量低于 15 L/min 时，保护效果差，焊缝氧化，飞溅增加。
**处理方法**：检查气体管路，确认流量 ≥ 18 L/min。

### 原因三：焊丝伸出长度过长
焊丝伸出超过 15mm 时，电弧稳定性下降。
**处理方法**：调整焊枪高度，伸出控制在 8-12mm。

## 安全注意事项

焊接作业必须佩戴：
- 焊接面罩（滤光片 ≥ 10#）
- 防火手套
- 防护服
- 通风设备
"""


@pytest.fixture()
def sample_docx_bytes():
    """生成最小可用 docx 字节（python-docx 必须可用时才返回）"""
    try:
        from docx import Document  # python-docx

        doc = Document()
        doc.add_heading("数控机床主轴维护手册", level=1)
        doc.add_heading("主轴异响故障", level=2)
        p = doc.add_paragraph("数控机床主轴出现异响时，常见原因包括：")
        doc.add_paragraph("1. 轴承磨损：检查主轴轴承游隙，超过 0.02mm 需更换")
        doc.add_paragraph("2. 润滑不足：润滑油变质或油路堵塞")
        doc.add_paragraph("3. 联轴器松动：检查螺栓紧固扭矩")
        doc.add_paragraph("处理方法：停车检查 → 拆主轴 → 测轴承 → 更换损坏部件")
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except ImportError:
        pytest.skip("python-docx 未安装，跳过 docx 测试")


@pytest.fixture()
def sample_pdf_bytes():
    """生成最小可用 PDF 字节（pypdf 可写但优先用 reportlab）"""
    try:
        # 优先用 reportlab 生成可解析的 PDF
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import A4

        buf = io.BytesIO()
        c = canvas.Canvas(buf, pagesize=A4)
        c.drawString(100, 750, "液压系统维修手册")
        c.drawString(100, 700, "液压泵噪声过大故障")
        c.drawString(100, 670, "原因分析：")
        c.drawString(100, 640, "1. 油液污染：过滤网堵塞导致吸油不畅")
        c.drawString(100, 610, "2. 轴承磨损：泵轴轴承寿命到期")
        c.drawString(100, 580, "3. 气穴现象：油箱液位过低")
        c.drawString(100, 550, "处理方案：清洗滤油器、更换液压油、检查吸油管路")
        c.save()
        return buf.getvalue()
    except ImportError:
        # 兜底：pypdf 写一个空 PDF
        try:
            from pypdf import PdfWriter
            writer = PdfWriter()
            writer.add_blank_page(width=595, height=842)
            buf = io.BytesIO()
            writer.write(buf)
            return buf.getvalue()
        except ImportError:
            pytest.skip("reportlab 和 pypdf 都未安装，跳过 PDF 测试")


# ============================================================
# 文件上传 helper
# ============================================================

def make_upload_file(filename: str, content: bytes, mime: str = "application/octet-stream"):
    """构造 multipart/form-data 上传的文件部分"""
    return (filename, content, mime)
